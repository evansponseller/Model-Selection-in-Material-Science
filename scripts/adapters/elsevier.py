"""Elsevier/Scopus publisher adapter."""

import io
import re
import time
import xml.etree.ElementTree as ET

import requests

from config import (
    ARTICLE_PII_URL,
    ARTICLE_URL,
    BOILERPLATE_SECTIONS,
    ELSEVIER_API_KEY,
    HTTP_REFERER,
    MAX_SCOPUS_RESULTS,
    RATE_LIMIT_PAUSE,
    SCOPUS_QUERY,
    SCOPUS_URL,
)
from adapters.base import PublisherAdapter

# ── Supplementary-material parsing/extraction ────────────────────────────────
# Elsevier lists supplementary files as <object ref="mmcN" ...>DOWNLOAD_URL</object>
# inside the full-text XML. Each has a mimetype and a content/object download URL.
_OBJECT_RE = re.compile(r"<object\b([^>]*)>([^<]+)</object>")


def _parse_supp_objects(xml_text: str) -> list[tuple[str, str, str]]:
    """Return (ref, mimetype, url) for every supplementary object (ref starts 'mmc')."""
    out: list[tuple[str, str, str]] = []
    for m in _OBJECT_RE.finditer(xml_text):
        attrs, url = m.group(1), m.group(2).strip()
        ref_m = re.search(r'\bref="([^"]+)"', attrs)
        mime_m = re.search(r'\bmimetype="([^"]+)"', attrs)
        if ref_m and ref_m.group(1).lower().startswith("mmc"):
            out.append((ref_m.group(1), (mime_m.group(1) if mime_m else "").lower(), url))
    return out


def _pdf_to_text(data: bytes) -> str:
    from pypdf import PdfReader
    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # noqa: BLE001 — supp parsing is best-effort
        print(f"  [elsevier] PDF parse failed: {exc}")
        return ""


def _docx_to_text(data: bytes) -> str:
    import docx
    try:
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001 — supp parsing is best-effort
        print(f"  [elsevier] DOCX parse failed: {exc}")
        return ""


def _extract_supp_text(mimetype: str, url: str, data: bytes) -> str:
    """Dispatch text extraction by mimetype/extension. Skips images, spreadsheets, zips."""
    u = url.lower()
    if "pdf" in mimetype or ".pdf" in u:
        return _pdf_to_text(data)
    if "word" in mimetype or "wordprocessing" in mimetype or ".docx" in u:
        return _docx_to_text(data)
    return ""  # unsupported supp type (xls, zip, image, video) — skip

# Elsevier XML namespaces
_NS = {
    "xocs":  "http://www.elsevier.com/xml/xocs/dtd",
    "ce":    "http://www.elsevier.com/xml/common/dtd",
    "dc":    "http://purl.org/dc/elements/1.1/",
    "prism": "http://prismstandard.org/namespaces/basic/2.0/",
    "tb":    "http://www.elsevier.com/xml/common/table/dtd",
}


def _get_with_retry(url: str, headers: dict, params: dict | None = None) -> dict:
    """GET request with simple retry on 429."""
    from config import MAX_RETRIES, RATE_LIMIT_RETRY_PAUSE
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            resp = requests.get(url, headers=headers, params=params, timeout=30)
            if resp.status_code == 429:
                wait = RATE_LIMIT_RETRY_PAUSE * attempt
                print(f"  [elsevier] Rate limited. Waiting {wait}s …")
                time.sleep(wait)
                continue
            resp.raise_for_status()
            return resp.json()
        except requests.RequestException as exc:
            from config import RATE_LIMIT_RETRY_PAUSE as pause
            wait = pause * attempt
            print(f"  [elsevier] Network error (attempt {attempt}): {exc}. Waiting {wait}s …")
            time.sleep(wait)
    return {}


class ElsevierAdapter(PublisherAdapter):
    """Adapter for Elsevier Scopus search + ScienceDirect full-text retrieval."""

    def search(self, max_results: int = MAX_SCOPUS_RESULTS) -> list[dict]:
        """Return papers from Scopus as {title, paper_id, doi, abstract, source} dicts."""
        if not ELSEVIER_API_KEY:
            print("  [elsevier] ELSEVIER_API_KEY not set — skipping Elsevier search.")
            return []

        headers = {"X-ELS-APIKey": ELSEVIER_API_KEY, "Accept": "application/json"}
        results: list[dict] = []
        start = 0
        page_size = 25

        while len(results) < max_results:
            params = {
                "query": SCOPUS_QUERY,
                "count": page_size,
                "start": start,
                "field": "dc:title,prism:doi,dc:description,prism:publicationName",
            }
            data = _get_with_retry(SCOPUS_URL, headers=headers, params=params)
            if start == 0:
                total = data.get("search-results", {}).get("opensearch:totalResults", "?")
                print(f"  Scopus total results: {total}")
            entries = data.get("search-results", {}).get("entry", [])
            if not entries:
                break
            for e in entries:
                doi = e.get("prism:doi", "").strip()
                if doi:
                    results.append({
                        "title": e.get("dc:title", "").strip(),
                        "paper_id": doi,
                        "doi": doi,
                        "abstract": e.get("dc:description", "").strip(),
                        "journal": e.get("prism:publicationName", "").strip(),
                        "source": "elsevier",
                    })
            print(f"  Fetched {len(results)} Scopus candidates so far …", end="\r")
            start += page_size
            if len(entries) < page_size:
                break
            time.sleep(RATE_LIMIT_PAUSE)
        print(f"  Fetched {len(results)} Scopus candidates total.      ")

        return results[:max_results]

    def fetch_full_text(self, paper: dict) -> dict | None:
        """Fetch full XML from Elsevier by DOI (or PII if doi missing)."""
        doi = paper.get("doi") or paper.get("paper_id", "")
        if not doi:
            return None
        url = ARTICLE_URL.format(doi=doi)
        return self._fetch_and_parse_xml(url)

    def fetch_full_text_by_pii(self, pii: str) -> dict | None:
        """Fetch full XML from Elsevier by PII."""
        url = ARTICLE_PII_URL.format(pii=pii)
        return self._fetch_and_parse_xml(url)

    def _fetch_and_parse_xml(self, url: str) -> dict | None:
        if not ELSEVIER_API_KEY:
            raise EnvironmentError("ELSEVIER_API_KEY is not set.")

        headers = {
            "X-ELS-APIKey": ELSEVIER_API_KEY,
            "Accept": "text/xml",
            "HTTP-Referer": HTTP_REFERER,
        }
        try:
            resp = requests.get(url, headers=headers, timeout=30)
            resp.raise_for_status()
        except requests.RequestException as exc:
            print(f"  [elsevier] Error fetching {url}: {exc}")
            return None

        try:
            root = ET.fromstring(resp.content)
        except ET.ParseError as exc:
            print(f"  [elsevier] XML parse error for {url}: {exc}")
            return None

        result = self._extract_sections(root)
        # Always capture intro text for classification, even though it's excluded from extraction
        if "intro_text" not in result:
            result["intro_text"] = self._extract_intro(root)
        return result

    def _extract_sections(self, root: ET.Element) -> dict:
        """Pull section titles + text from Elsevier XML into a dict."""
        sections: dict[str, str] = {}
        tables: list[str] = []

        for section in root.iter():
            tag = section.tag.split("}")[-1] if "}" in section.tag else section.tag
            if tag == "section":
                title_el = section.find(
                    ".//{http://www.elsevier.com/xml/common/dtd}section-title"
                )
                title = (
                    title_el.text.strip().lower()
                    if title_el is not None and title_el.text
                    else "unknown"
                )
                text = " ".join(t.strip() for t in section.itertext() if t.strip())
                if title not in BOILERPLATE_SECTIONS:
                    sections[title] = text
            elif tag == "table":
                table_text = " ".join(t.strip() for t in section.itertext() if t.strip())
                if table_text:
                    tables.append(table_text)

        return {"sections": sections, "tables": tables}

    def _extract_intro(self, root: ET.Element) -> str:
        """Extract introduction text specifically for classification use."""
        for section in root.iter():
            tag = section.tag.split("}")[-1] if "}" in section.tag else section.tag
            if tag == "section":
                title_el = section.find(
                    ".//{http://www.elsevier.com/xml/common/dtd}section-title"
                )
                title = (
                    title_el.text.strip().lower()
                    if title_el is not None and title_el.text
                    else ""
                )
                if title == "introduction":
                    return " ".join(t.strip() for t in section.itertext() if t.strip())
        return ""

    # ── Supplementary materials ──────────────────────────────────────────────

    def fetch_supplementary_text(self, doi: str) -> str:
        """
        Download and extract text from a paper's supplementary files (mmcN objects).
        Returns the combined text of all PDF/DOCX supplements, labelled by ref.
        Returns "" if there are none or none are text-extractable.
        """
        if not ELSEVIER_API_KEY:
            raise EnvironmentError("ELSEVIER_API_KEY is not set.")

        headers = {
            "X-ELS-APIKey": ELSEVIER_API_KEY,
            "Accept": "text/xml",
            "HTTP-Referer": HTTP_REFERER,
        }
        try:
            resp = requests.get(ARTICLE_URL.format(doi=doi), headers=headers, timeout=30)
            resp.raise_for_status()
        except requests.RequestException as exc:
            print(f"  [elsevier] Error fetching XML for supp ({doi}): {exc}")
            return ""

        objects = _parse_supp_objects(resp.text)
        if not objects:
            return ""

        chunks: list[str] = []
        for ref, mimetype, url in objects:
            data = self._download_object(url)
            if not data:
                continue
            text = _extract_supp_text(mimetype, url, data)
            if text.strip():
                chunks.append(f"[{ref}]\n{text.strip()}")
            else:
                print(f"  [elsevier] {ref}: unsupported/empty supp type ({mimetype})")
            time.sleep(RATE_LIMIT_PAUSE)

        return "\n\n".join(chunks)

    def _download_object(self, url: str) -> bytes | None:
        """Download a single Elsevier object (supplementary file) as bytes."""
        try:
            resp = requests.get(
                url, headers={"X-ELS-APIKey": ELSEVIER_API_KEY}, timeout=60
            )
            resp.raise_for_status()
            return resp.content
        except requests.RequestException as exc:
            print(f"  [elsevier] Error downloading object {url}: {exc}")
            return None
